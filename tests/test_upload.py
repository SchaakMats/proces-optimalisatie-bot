import io
import pytest


def test_extract_txt():
    from app import extract_text
    content = b"Dit is een testdocument.\nMet twee regels."
    result = extract_text(content, "test.txt")
    assert "testdocument" in result
    assert "twee regels" in result


def test_extract_txt_encoding():
    from app import extract_text
    content = "Bedrijf: Café & Zonen".encode("utf-8")
    result = extract_text(content, "test.txt")
    assert "Café" in result


def test_extract_pdf():
    from app import extract_text
    import pypdf
    writer = pypdf.PdfWriter()
    writer.add_blank_page(width=200, height=200)
    buf = io.BytesIO()
    writer.write(buf)
    result = extract_text(buf.getvalue(), "test.pdf")
    assert isinstance(result, str)


def test_extract_docx():
    from app import extract_text
    import docx
    doc = docx.Document()
    doc.add_paragraph("Procesoptimalisatie test")
    doc.add_paragraph("Tweede alinea")
    buf = io.BytesIO()
    doc.save(buf)
    result = extract_text(buf.getvalue(), "test.docx")
    assert "Procesoptimalisatie test" in result
    assert "Tweede alinea" in result


def test_extract_unsupported_type():
    from app import extract_text
    with pytest.raises(ValueError, match="Niet ondersteund"):
        extract_text(b"data", "test.xlsx")


def test_upload_txt_returns_text(client):
    data = {"file": (io.BytesIO(b"Testinhoud voor upload"), "test.txt")}
    resp = client.post("/api/upload", data=data, content_type="multipart/form-data")
    assert resp.status_code == 200
    result = resp.get_json()
    assert result["filename"] == "test.txt"
    assert "Testinhoud" in result["text"]
    assert result["chars"] > 0


def test_upload_no_file_returns_400(client):
    resp = client.post("/api/upload", data={}, content_type="multipart/form-data")
    assert resp.status_code == 400


def test_upload_unsupported_type_returns_400(client):
    data = {"file": (io.BytesIO(b"data"), "test.xlsx")}
    resp = client.post("/api/upload", data=data, content_type="multipart/form-data")
    assert resp.status_code == 400


def test_upload_empty_file_returns_400(client):
    data = {"file": (io.BytesIO(b"   "), "leeg.txt")}
    resp = client.post("/api/upload", data=data, content_type="multipart/form-data")
    assert resp.status_code == 400


def test_upload_docx_returns_text(client):
    import docx
    doc = docx.Document()
    doc.add_paragraph("Testbedrijf intake document")
    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    data = {"file": (buf, "rapport.docx")}
    resp = client.post("/api/upload", data=data, content_type="multipart/form-data")
    assert resp.status_code == 200
    result = resp.get_json()
    assert "Testbedrijf" in result["text"]
